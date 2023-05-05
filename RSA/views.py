"""
Routes and views for the flask application.
Author: Dr. Reza Dilmaghani
"""

from datetime import datetime
from flask import render_template, render_template, request, redirect, jsonify, make_response, send_from_directory, send_file
from RSA import app
import PyPDF2
import utils as ut
import pandas as pd
import os
import numpy as np
import re
import extract_msg
import docx
import spacy
import nltk
import itertools
from nltk.tokenize import sent_tokenize

app.config['UPLOAD_FOLDER'] = './RSA/uploads/'
app.config['DOWNLOAD_FOLDER'] = './RSA/downloads/'
app.config['ATTACH_FOLDER'] = './RSA/Attachments/'
uk_postcode_regex = r'[A-Z]{1,2}\d[A-Z\d]? ?\d[A-Z]{2}'

# Load the English model
nlp = spacy.load("en_core_web_sm")

###########################
### List of Functions
###########################
# Load the docx file
def read_docx_file(filepath):
    doc = docx.Document(filepath)
    doc_content = ""
    for para in doc.paragraphs:
        doc_content = doc_content + " " + para.text +"."
    doc_content = doc_content.strip() + "." + '\n'               
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                doc_content = doc_content + " " + cell.text +"."
    return doc_content
    

# Load the msg file
def read_msg_file(msg_file_path):
    # Create a Message object
    msg = extract_msg.Message(msg_file_path)

    # Extract the information you need from the .msg file
    subject = msg.subject
    sender = msg.sender
    msg_date = msg.date
    body = msg.body
    
    sender_name = re.sub(r'<.*?>', '', sender).strip()
    sender_email = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', sender).group().strip()
    
    # Extract the HTML body of the message
    html_body = msg.htmlBody.decode('utf-8', errors='ignore')
    
    # Find the signature in the HTML body
    signature_regex = r'<div.*?class="?Signature"?[^>]*>(.*?)</div>'
    signature_match = re.search(signature_regex, html_body, re.IGNORECASE | re.DOTALL)
    
    signature =''
    if signature_match:
        # Extract the signature text
        signature = signature_match.group(1)
        # Remove any remaining HTML tags from the signature
        signature = re.sub('<.*?>', '', signature)
    
    # Loop through the attachments in the message
    att_names = []
    for attachment in msg.attachments:
        # Check if the attachment is a Word document
        if attachment.longFilename.lower().endswith(".docx") or attachment.longFilename.lower().endswith(".doc"):
            # Save the attachment to the specified directory
            filename = os.path.join(app.config['ATTACH_FOLDER'], attachment.longFilename)
            os.chdir(app.config['ATTACH_FOLDER'])
            attachment.save()
            att_names.append(attachment.longFilename)
    msg.close()
    return subject, sender, sender_name, sender_email, msg_date, body, signature, att_names


def nlp_add_msg(clean_text):
    doc = nlp(clean_text)
    address_components = ''
    for i, sent in enumerate(doc.sents):    
        postcode = re.findall(uk_postcode_regex, sent.text)
        if postcode:
            for j, ent in enumerate(sent.ents):                                                                     
                if ent.label_ == "WORK_OF_ART":
                    if j < len(sent.ents)-1: 
                        if sent.ents[j+1].label_ == "GPE" or sent.ents[j+1].label_ == "FAC":
                            address_components  += ent.text + ", "              

                if ent.label_ == "CARDINAL":
                    if j < len(sent.ents)-1: 
                        if sent.ents[j+1].label_ == "GPE" or sent.ents[j+1].label_ == "FAC":
                            address_components  += ent.text + " "
                
                if ent.label_ == 'FAC':                
                    address_components  += ent.text + ", " 
                
                if ent.label_ == 'GPE':
                    # Extract the town/city
                    if ent.text not in postcode[0]: 
                        address_components  += ent.text + ", " 
            
            address_components += postcode[0]
            break
    return address_components, postcode[0]
    
def nlp_add_msg_client(clean_text):
    doc = nlp(clean_text)
    address_components = ''
    for i, sent in enumerate(doc.sents):    
        postcode = re.findall(uk_postcode_regex, sent.text)
        if postcode:
            for j, ent in enumerate(sent.ents):                                                                     
                if ent.label_ == "WORK_OF_ART":
                    if j < len(sent.ents)-1: 
                        if sent.ents[j+1].label_ == "GPE" or sent.ents[j+1].label_ == "FAC":
                            address_components  += ent.text + ", "              

                if ent.label_ == "CARDINAL":
                    if j < len(sent.ents)-1: 
                        if sent.ents[j+1].label_ == "GPE" or sent.ents[j+1].label_ == "FAC":
                            address_components  += ent.text + " "
                
                if ent.label_ == 'FAC':                
                    address_components  += ent.text + ", " 
                
                if ent.label_ == 'GPE':
                    # Extract the town/city
                    if ent.text not in postcode[0]: 
                        address_components  += ent.text + ", " 
            
            address_components += postcode[0]
            break
    return address_components, postcode[0]


    
def nlp_answer(doc_input):   
    questions = [
                   "What is the name of broker?"
                ]
        
    sentences = sent_tokenize(doc_input)
    nlp_answers = ''
    # Answer Extraction
    for question in questions:
        doc = nlp(question)
        for sentence in sentences:
            doc = nlp(sentence)
            if 'brokers' in sentence.lower():
                for token in doc:
                    print ("-------------")
                    print (token)
                    print ("-------------")
                    if token.head.text.lower() == 'brokers'  and len(list(token.children)) == 1:
                        nlp_answers = str(next(itertools.islice(token.children, 1))) + " " + str(token.text) + " " + str(token.head.text)
                        break;
    return nlp_answers
    

def client_details(doc_input):   
    doc = nlp(doc_input)
    client_name_pos = doc_input.lower().find("client name") # find the position of "client name" in the sentence

    closest_entity = None
    min_distance = float("inf")

    for ent in doc.ents:
        if ent.label_ in ["ORG"]:
            distance = abs(ent.start_char - client_name_pos) 
            if distance < min_distance:
                closest_entity = ent
                min_distance = distance
    return closest_entity.text
    
    

###########################
### Run the program
###########################

@app.route('/')
@app.route('/home')
def home():
    """Renders the home page."""
    return render_template(
        'index.html',
        title='Home Page',
        year=datetime.now().year,
    )

@app.route('/_RSA', methods=['GET', 'POST'])
def _RSA():
    #ut.empty_upload_download(app.config['UPLOAD_FOLDER'], app.config['DOWNLOAD_FOLDER'])
    rsa_file = request.files['rsaFile']  
    filename_rsa, rsa_filetype = ut.save_file(app.config['UPLOAD_FOLDER'], rsa_file)     
    
    if rsa_filetype =="msg":
        subject, sender, sender_name, sender_email, msg_date, msg_body, signature, att_names = read_msg_file(app.config['UPLOAD_FOLDER'] + filename_rsa)
        
    if att_names:
        msg_doc_content = read_docx_file(os.getcwd() + "/" + att_names[0])
        rsa_address, postcode = nlp_add_msg(msg_body) 
        broker_name = nlp_answer(msg_doc_content)
    
 
    
    Broker_Details = pd.DataFrame(columns=['Broker Name', 'Broker Contact Name', 'Broker Contact Email', 'Date', 'Address', 'Post Code', 'Attachment'])
    
    Broker_Details = Broker_Details.append({
                                            'Broker Name' : broker_name, 
                                            'Broker Contact Name': sender_name, 
                                            'Broker Contact Email': sender_email, 
                                            'Date': msg_date, 
                                            'Address': rsa_address, 
                                            'Post Code': postcode, 
                                            'Attachment': att_names[0]
                                            }, ignore_index=True) 
    
    Broker_Details = Broker_Details.transpose()
    broker_doc= att_names[0]     

    #################################################
    ####### Client Details
    #################################################   
    Client_Dtl = pd.DataFrame(columns=['Client Name', 'Activity', 'Insurer', 'Web page', 'Address'])
    Client_info = client_details(msg_doc_content)
    
    Client_Dtl = Client_Dtl.append({
                                            'Client Name' : Client_info,
                                            'Activity': '',
                                            'Insurer': '',    
                                            'Web page': '',
                                            'Address': ''                                            
                                            }, ignore_index=True) 
    Client_Dtl = Client_Dtl.transpose()

    #################################################
    ####### Reutrn the results
    #################################################  
    html_data = "1"    
    return jsonify(data=html_data, myt=Broker_Details.to_html(classes='table table-striped table-bordered" name = "mytablediv" id = "mytablediv"'                                                                                                                        
                                                              , index=True, header= False,  escape=False, border=0),
                                   myt2=Client_Dtl.to_html(classes='table table-striped table-bordered" name = "mytablediv2" id = "mytablediv2"'
                                                               , index=True, header= False,  escape=False, border=0)                        
                                                               , broker_doc=broker_doc                                                             
                   )
                          
                              
                                                                        
@app.route('/get_file', methods=['GET', 'POST'])
def get_file():
    """Download a file."""    
    resultedFile = request.args['resultedFile']
        
    finalFile = "Attachments/" + resultedFile
    return send_file(finalFile, as_attachment=True)
    
    
##########################################################################
